import zipfile
from copy import deepcopy
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional, Tuple

from deepdiff import DeepDiff
from docx import Document as OpenDocument
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml.text.run import CT_R
from docx.oxml.xmlchemy import BaseOxmlElement
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from lxml.etree import QName

from backend.my_docx.my_docx import Docx
from logger import f, logger
from utils.image import images_equal
from vars import PATH_TMP

# ------------------- Getter -------------------


def _get_default_paragraph_style(doc):
    for style in doc.styles:
        if style.type != WD_STYLE_TYPE.PARAGRAPH:
            continue

        el = style._element
        if el.get(qn("w:default")) == "1":
            return style

    return None


def _safe_normal_style(doc: Docx):
    try:
        return doc.styles["Normal"]
    except KeyError:
        return _get_default_paragraph_style(doc)


def _get_run_background(run: Run):
    r = run._r
    shd = r.xpath("./w:rPr/w:shd")
    if not shd:
        return None
    return shd[0].get(qn("w:fill"))  # hex color or "auto"


def _get_paragraph_background(paragraph: Paragraph):

    p = paragraph._p
    shd = p.xpath("./w:pPr/w:shd")
    if not shd:
        return None
    return shd[0].get(qn("w:fill"))


def _get_cell_background(cell):
    """
    Returns the cell background color as a hex string (e.g. 'FFFF00'),
    'auto', or None if not set.
    """
    tc = cell._tc
    shd = tc.xpath("./w:tcPr/w:shd")
    if not shd:
        return None
    return shd[0].get(qn("w:fill"))


def _get_effective_background(run):
    # 1. Run shading
    bg = _get_run_background(run)
    if bg:
        return bg

    # 2. Paragraph shading
    bg = _get_paragraph_background(run._parent)
    if bg:
        return bg

    # 3. Highlight (marker)
    if run.font.highlight_color:
        return str(run.font.highlight_color)

    return None


def _normalize_style(el) -> Optional[str]:
    if el.style is None:
        return "Normal"

    if el.style.name in ["Default Paragraph Font", "Normal", "_Normal"]:
        return "Normal"

    return el.style.name


def _effective_font_name(doc: Docx, p: Paragraph, r: CT_R) -> Optional[str]:
    # 1. Run-level
    rPr = r.find(qn("w:rPr"))
    if rPr is not None and rPr.find(qn("w:rFonts")) is not None:
        rPr.find(qn("w:rFonts")).get(qn("w:ascii"))

    # 2. Paragraph style
    p_style = p.style
    if p_style and p_style.font.name:
        return p_style.font.name

    # 3. Document defaults (Normal)
    normal = _safe_normal_style(doc)
    if normal and normal.font.name:
        return normal.font.name

    # 4. Theme-based (unknown concrete font)
    return None  # resolved by Word at render time


def _effective_font_size(doc: Docx, p: Paragraph, r: CT_R) -> Optional[str]:
    # 1. Run-level
    rPr = r.find(qn("w:rPr"))
    if rPr is not None and rPr.find(qn("w:sz")) is not None:
        return int(rPr.find(qn("w:sz")).get(qn("w:val"))) / 2

    # 2. Paragraph style
    if p.style and p.style.font.size:
        return p.style.font.size.pt

    # 3. Normal style
    normal = _safe_normal_style(doc)
    if normal and normal.font.size:
        return normal.font.size.pt

    # Word default ≈ 11pt (but technically theme-based)
    return 11.0


def _effective_font_color(doc: Docx, p: Paragraph, r: CT_R) -> Optional[str]:

    # 1. Explicit RGB
    rPr = r.find(qn("w:rPr"))
    if rPr is not None and rPr.find(qn("w:color")) is not None:
        return rPr.find(qn("w:color")).get(qn("w:val"))

    # 3. Paragraph style
    p_style = p.style
    if p_style and p_style.font.color.rgb:
        return tuple(p_style.font.color.rgb)

    # 4. Normal style
    normal = _safe_normal_style(doc)
    if normal and normal.font.color.rgb:
        return tuple(normal.font.color.rgb)

    # Default: automatic (usually black)
    return "000000"


def extract_text_from_run_xml(r: CT_R) -> str:
    return "".join(t.text or "" for t in r.findall(qn("w:t")))


def _extract_run_from_xml(doc: Docx, p: Paragraph, r: CT_R) -> Dict[str, Any]:
    assert _is_run(r)
    rPr = r.find(qn("w:rPr"))

    def get_bool(tag):
        if rPr is None:
            return None
        el = rPr.find(qn(f"w:{tag}"))
        if el is None:
            return None
        val = el.get(qn("w:val"))
        return val != "0"

    def get_val(tag):
        if rPr is None:
            return None
        el = rPr.find(qn(f"w:{tag}"))
        if el is None:
            return None
        return el.get(qn("w:val"))

    return {
        "text": extract_text_from_run_xml(r),
        "bold": get_bool("b"),
        "italic": get_bool("i"),
        "underline": get_val("u"),
        "font_name": _effective_font_name(doc, p, r),
        "font_size": _effective_font_size(doc, p, r),
        "color_rgb": _effective_font_color(doc, p, r),
        "highlight": (
            rPr.find(qn("w:highlight")).get(qn("w:val"))
            if rPr is not None and rPr.find(qn("w:highlight")) is not None
            else None
        ),
        "background": (
            rPr.find(qn("w:shd")).get(qn("w:fill"))
            if rPr is not None and rPr.find(qn("w:shd")) is not None
            else None
        ),
    }


def _localname(el: BaseOxmlElement) -> str:
    return QName(el).localname


# ------------------- Image -------------------


def _extract_image(path_docx: Path, image_rel_path: Path, path_folder_output) -> Path:
    with zipfile.ZipFile(path_docx) as z:
        for f in z.namelist():
            if f.endswith(image_rel_path):
                z.extract(f, path_folder_output)
                return path_folder_output / f


def _get_images_rel_path(doc: Docx, p: Paragraph) -> List[Path]:

    images: List[Path] = []
    for run in p.runs:
        drawings = run._element.xpath(".//w:drawing")
        if drawings:
            blips = run._element.xpath(".//a:blip")
            for blip in blips:
                rId = blip.get(
                    "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                )
                rel_path = doc.part._rels[rId].target_ref
                images.append(rel_path)

    return images


def _get_images(doc: Docx, p: Paragraph) -> List[Path]:
    return [
        _extract_image(doc.path, image_rel_path=rel_path, path_folder_output=PATH_TMP)
        for rel_path in _get_images_rel_path(doc, p)
    ]


# ------------------- Predicate -------------------


def _is_run(xml_element) -> bool:
    return _localname(xml_element) == "r"


def _is_hyperlink(xml_element) -> bool:
    return _localname(xml_element) == "hyperlink"


def _run_signature(doc: Docx, p: Paragraph, run: CT_R):
    d = _extract_run_from_xml(doc, p, run)
    d.pop("text")
    return d


def _same_format(doc: Docx, p: Paragraph, a: CT_R, b: CT_R):
    return _run_signature(doc, p, a) == _run_signature(doc, p, b)


# ------------------- Merge -------------------


def is_protected_run(run: Run) -> bool:
    return _is_protected_run_xml(run._r)


def _is_protected_run_xml(r: CT_R) -> bool:
    return bool(
        r.xpath(".//w:drawing")
        or r.xpath(".//w:fldChar")
        or r.xpath(".//w:instrText")
        or r.xpath(".//w:bookmarkStart")
        or r.xpath(".//w:bookmarkEnd")
    )


def _can_merge(
    doc: Docx, p: Paragraph, el1: BaseOxmlElement, el2: BaseOxmlElement
) -> bool:
    return (
        _is_run(el1)
        and _is_run(el2)
        and _same_format(doc, p, el1, el2)
        and not _is_protected_run_xml(el1)
        and not _is_protected_run_xml(el2)
    )


def normalize_runs(doc: Docx, p: Paragraph, inplace: bool) -> Paragraph:

    # inplace
    if not inplace:
        doc = OpenDocument()
        p = Paragraph(deepcopy(p._p), parent=doc)

    # normal

    # normalize runs : merge same format together
    elements = list(p._p)
    i = 0

    while i < len(elements) - 1:
        current = elements[i]
        next_el = elements[i + 1]

        # Skip
        if not _can_merge(doc, p, current, next_el):
            i += 1
            continue

        # Merge text
        current.text += next_el.text

        # Remove next run XML
        p._p.remove(next_el)

        # Update list after removal
        elements.pop(i + 1)

        # Do NOT increment i — keep merging forward

    return p


# ------------------- Format -------------------


def _tip_paragraph(p: Paragraph) -> str:
    return f(text=f"{p.text[:30]}...") if p.text else ""


# ------------------- Equality -------------------


def _extract_paragraph(doc: Docx, p: Paragraph, normalize: bool) -> Dict[str, Any]:

    text = p.text

    p2 = normalize_runs(doc, p, inplace=False) if normalize else p
    runs = [_extract_run_from_xml(doc, p, el) for el in p2._p if _is_run(el)]

    return {
        "type": "paragraph",
        "style": _normalize_style(p),
        "alignment": p.alignment,
        "runs": runs,
        "id": text,
    }


def _extract_doc_structure(doc: Docx, normalize: bool) -> List[Dict[str, Any]]:
    blocks: List[Dict[str, Any]] = []

    # parapgraph
    for p in doc.paragraphs:
        blocks.append(_extract_paragraph(doc, p, normalize))

    # table
    for table in doc.tables:
        table_data = []
        text = ""
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_paragraphs = [
                    _extract_paragraph(doc, p, normalize) for p in cell.paragraphs
                ]
                row_data.append(
                    {
                        "type": "cell",
                        "background": _get_cell_background(cell),
                        "content": cell_paragraphs,
                    }
                )
                text += cell.text + "\n"

            table_data.append(row_data)

        blocks.append({"type": "table", "content": table_data, "id": text[:30]})

    return blocks


def _compare_structures(a: list, b: list) -> bool:
    if len(a) != len(b):
        logger.info(f"Different number of blocks : {len(a)} != {len(b)}")
        return False

    for i, (block_a, block_b) in enumerate(zip(a, b)):
        if block_a["type"] != block_b["type"]:
            logger.info(
                f"Block {i}: different types : {block_a['type']} != {block_b['type']}\n"
                + f(id1=block_a["id"], id2=block_b["id"])
            )
            return False

        diff = DeepDiff(block_a, block_b)
        if diff != {}:
            logger.info(
                f"Block {i}: content mismatch: diff:\n{diff}\n"
                + f(id1=block_a["id"], id2=block_b["id"])
            )
            return False

    return True


def _docx_image_equals(d1: Docx, d2: Docx) -> bool:

    for p1, p2 in zip(d1.paragraphs, d2.paragraphs):
        # get images in order
        images1 = _get_images(d1, p1)
        images2 = _get_images(d2, p2)

        if len(images1) != len(images2):
            logger.info(
                f"Docx equality nb images {_tip_paragraph(p1)} : {len(images1)} != {len(images2)}"
            )
            return False

        for i1, i2 in zip(images1, images2):
            if not images_equal(i1, i2):
                logger.info(f"Docx equality image {_tip_paragraph(p1)}")
                return False

        return True


def paragraph_equals(
    doc: Docx, p1: Paragraph, p2: Paragraph, normalize: bool = True
) -> bool:
    """Not comparing images"""

    e1 = _extract_paragraph(doc, p1, normalize=normalize)
    e2 = _extract_paragraph(doc, p2, normalize=normalize)

    return _compare_structures([e1], [e2])


def docx_equals(d1: Docx, d2: Docx, normalize: bool = True) -> bool:

    # paragraphs and tables
    struct1 = _extract_doc_structure(d1, normalize)
    struct2 = _extract_doc_structure(d2, normalize)

    if not _compare_structures(struct1, struct2):
        return False

    # image
    if not _docx_image_equals(d1, d2):
        return False

    return True


# ------------------- Modifier -------------------


def replace_text_paragraphs(
    doc: Docx,
    paragraphs: List[Paragraph],
    replace_text: Callable[[str], Tuple[str, int]],
) -> int:

    nb_changes = 0

    for p in paragraphs:
        normalize_runs(doc, p, inplace=True)

        # compute new runs
        for run in p.runs:

            if is_protected_run(run):
                continue

            changed_text, nb_new_changes = replace_text(run.text)

            run.text = changed_text
            nb_changes += nb_new_changes

    return nb_changes


def duplicate_paragraphs(
    doc: Docx, start_idx_paragraph: int, end_idx_paragraph: int, n: int
) -> int:
    """
    Returns:
        int: number of paragraphs added
    """

    assert start_idx_paragraph > 0

    start_p = doc.paragraphs[start_idx_paragraph - 1]._p

    # Paragraphs to duplicate
    block = doc.paragraphs[start_idx_paragraph : end_idx_paragraph + 1]

    # Insert duplicated block after 'start'
    for _ in range(n):
        for p in reversed(block):
            new_p = deepcopy(p._p)
            start_p.addnext(new_p)

    return n * len(block)


def remove_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._p
    parent = p.getparent()
    parent.remove(p)


# ------------------- Main -------------------

if __name__ == "__main__":

    from vars import PATH_TEST_DOCS, PATH_TEST_DOCS_TESTSUITE, PATH_TMP

    pass

    # p1 = (
    #     PATH_TEST_DOCS
    #     / "test_generation/TJ NAP1 - Visio administrative & Convocation modele.docx"
    # )
    # p1 = PATH_TEST_DOCS_TESTSUITE / "generation" / "docx" / "image_between_text.docx"
    p1 = PATH_TEST_DOCS_TESTSUITE / "docx" / "merge" / "hyperlink_between_text.docx"
    # p1 = PATH_TEST_DOCS_TESTSUITE / "docx" / "not_equals" / "bold.docx"

    doc = Docx(p1)

    # for p in doc.paragraphs[:20]:
    #     for run in p.runs:
    #         print(f"'{run.text}'")
    #         print("protected :", _is_hyperlink_run(run))
    # struct = _extract_doc_structure(doc)
    # print(struct[:20])

    for p in doc.paragraphs:
        for child in p._p:
            print(child.text, _localname(child), type(child))
            if _localname(child) == "hyperlink":
                print("hyperlink")
            elif _localname(child) == "r":
                print(f"run : '{child.text}'")
                print(_extract_run_from_xml(child))
            else:
                print("other :", child.tag)
