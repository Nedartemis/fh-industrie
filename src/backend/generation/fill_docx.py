from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path
from typing import List

from docx.oxml.text.run import CT_R
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from backend.generation.list.fill_list_helper import (
    RowInstruction,
    build_fullname_info,
    is_instruction,
    preprocess_instructions,
)
from backend.generation.list.fill_table_list import (
    is_the_table_a_table_list,
    replace_table_list,
)
from backend.generation.replace_text import build_replace_text
from backend.info_struct import InfoValues
from backend.my_docx.docx_helper import (
    duplicate_paragraphs,
    extract_text_from_run_xml,
    remove_paragraph,
    replace_text_paragraphs,
)
from backend.my_docx.docx_table import DocxTable
from backend.my_docx.my_docx import Docx
from logger import f, logger

# ------------------- Public Method -------------------


def fill_template_docx(
    template_path: Path, infos: InfoValues, path_output: Path
) -> int:

    doc = Docx(template_path)

    # without table

    # ind
    logger.debug("Replace independant infos")
    nb_changes = _replace_text_paragraphs_inds(doc, doc.paragraphs, infos)
    # list
    logger.debug("Replace lists infos without tables")
    nb_changes += _fill_list_without_table(doc, infos)

    # table
    logger.debug("Replace lists infos inside tables")
    nb_changes += _fill_tables(doc, infos)

    doc.save(path_output)

    return nb_changes


# ------------------- Private Method -------------------


def _replace_text_paragraphs_inds(
    doc: Docx, paragraphs: List[Paragraph], infos: InfoValues
):
    return replace_text_paragraphs(
        doc, paragraphs, build_replace_text(infos.independant_infos)
    )


def _fill_tables(doc: Docx, infos: InfoValues) -> int:

    nb_changes = 0

    # ind
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                nb_changes += _replace_text_paragraphs_inds(doc, cell.paragraphs, infos)

    # list
    for table in doc.tables:
        _fill_table_list(DocxTable(doc, table), infos)

    return nb_changes


def _fill_table_list(table: DocxTable, infos: InfoValues) -> int:

    if not is_the_table_a_table_list(table):
        return 0

    nb_changes = replace_table_list(table, infos)

    table.remove_column(col=1)

    return nb_changes


@dataclass
class RunWrapper:
    idx_paragraph: int
    run: CT_R

    def __repr__(self):
        return f"Run(idx={self.idx_paragraph}, text='{extract_text_from_run_xml(self.run)}')"


def _fill_list_without_table(doc: Docx, infos: InfoValues) -> int:

    childs: List[CT_R] = [child for p in doc.paragraphs for child in p._p]
    print(len(childs), [extract_text_from_run_xml(child) for child in childs])

    # find beginning and end
    row_instructions: List[RowInstruction[RunWrapper]] = [
        RowInstruction(
            text=text,
            tracability=RunWrapper(idx_paragraph=idx_paragraph, run=child),
        )
        for idx_paragraph, p in enumerate(doc.paragraphs)
        for child in p._p
        if (text := extract_text_from_run_xml(child))
        if is_instruction(text)
    ]

    # preprocess
    list_instructions = preprocess_instructions(row_instructions)

    # fill
    nb_changes = 0
    for instr in list_instructions:

        list_info = infos.list_infos.get(instr.first_name)
        if list_info is None:
            logger.info(f"{instr.first_name} not in infos")
            continue

        logger.debug(f"infos : {list_info}")

        nb_paragraphs_each_block = instr.end.idx_paragraph - (
            instr.start.idx_paragraph + 1
        )

        # duplicate paragraphs
        nb_paragraphs_added = duplicate_paragraphs(
            doc,
            start_idx_paragraph=instr.start.idx_paragraph + 1,
            end_idx_paragraph=instr.end.idx_paragraph - 1,
            n=len(list_info) - 1,
        )

        # update start and end
        for instr_other in list_instructions:
            if instr_other.start.idx_paragraph > instr.start.idx_paragraph:
                instr_other.start.idx_paragraph += nb_paragraphs_added

            if instr_other.end.idx_paragraph > instr.start.idx_paragraph:
                instr_other.end.idx_paragraph += nb_paragraphs_added

        # replace

        for idx, infos_one_element in enumerate(list_info):

            # build infos name
            infos_one_element = {
                build_fullname_info(instr.first_name, sub_name): value
                for sub_name, value in infos_one_element.items()
                if value is not None
            }

            # replace
            start_idx = instr.start.idx_paragraph + 1 + nb_paragraphs_each_block * idx
            end_idx = (
                instr.start.idx_paragraph + 1 + nb_paragraphs_each_block * (idx + 1)
            )
            nb_changes += replace_text_paragraphs(
                doc=doc,
                paragraphs=doc.paragraphs[start_idx:end_idx],
                replace_text=build_replace_text(infos_one_element),
            )

    # remove instructions
    paragraphs_to_remove = [
        doc.paragraphs[idx]
        for instr in list_instructions
        for idx in (instr.start.idx_paragraph, instr.end.idx_paragraph)
        if instr.first_name in infos.list_infos
    ]
    for p in paragraphs_to_remove:
        remove_paragraph(p)

    return nb_changes
