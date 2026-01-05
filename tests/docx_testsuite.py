import os
from pathlib import Path
from typing import Dict, Tuple

import pytest
from docx import Document as OpenDocument
from docx.table import Table
from docx.text.paragraph import Paragraph
from helper_testsuite import biv, wrapper_test_good, wrapper_try

from backend.generation.replace_text import build_replace_text
from backend.my_docx.docx_helper import docx_equals, normalize_runs, paragraph_equals
from backend.my_docx.docx_table import DocxTable
from backend.my_docx.my_docx import Docx
from logs_label import DuplicatesNameAfterHarmonization
from vars import PATH_TEST_DOCS_TESTSUITE

# ------------------- Equals -------------------


filenames_equals = [
    filename
    for filename in os.listdir(PATH_TEST_DOCS_TESTSUITE / "docx/not_equals")
    if not filename.endswith("not.docx")
]


@pytest.mark.parametrize(("filename"), filenames_equals)
def test_not_equals(filename: str):
    path_folder = PATH_TEST_DOCS_TESTSUITE / "docx/not_equals"
    path1 = path_folder / filename
    path2 = path_folder / f"{Path(filename).stem}_not.docx"

    def runnable():
        d1 = Docx(path1)
        d2 = Docx(path2)
        assert not docx_equals(d1, d2)

    wrapper_test_good(runnable=runnable)


@pytest.mark.parametrize(("filename"), filenames_equals)
def test_equals(filename: str):
    path_folder = PATH_TEST_DOCS_TESTSUITE / "docx/not_equals"
    path1 = path_folder / f"{filename}"

    def runnable():
        d1 = Docx(path1)
        d2 = Docx(path1)
        assert docx_equals(d1, d2)

    wrapper_test_good(runnable=runnable)


@pytest.mark.parametrize(("filename1", "filename2"), [("split_run", "split_run_not")])
def test_equals(filename1: str, filename2: str):
    path_folder = PATH_TEST_DOCS_TESTSUITE / "docx/equals"
    path1 = path_folder / f"{filename1}.docx"
    path2 = path_folder / f"{filename2}.docx"

    def runnable():
        d1 = Docx(path1)
        d2 = Docx(path2)
        assert docx_equals(d1, d2)

    wrapper_test_good(runnable=runnable)


# ------------------- Merge -------------------


def _equals_merge_not_in_place(d1: Docx, to_merge: Docx) -> bool:
    for p, p_to_merge in zip(d1.paragraphs, to_merge.paragraphs):
        p_merged = normalize_runs(d1, p_to_merge, inplace=False)
        assert paragraph_equals(d1, p, p_merged)


def _equals_merge_in_place(d1: Docx, merged: Docx) -> bool:

    for p in merged.paragraphs:
        normalize_runs(d1, p, inplace=True)

    assert docx_equals(d1, merged, normalize=False)


@pytest.mark.parametrize(("filename"), ["mix_style"])
def test_do_not_merge(filename: str):
    path_folder = PATH_TEST_DOCS_TESTSUITE / "docx/merge"
    path1 = path_folder / f"{filename}.docx"

    def runnable():
        d1 = Docx(path1)
        d2 = Docx(path1)

        _equals_merge_not_in_place(d1, d2)
        _equals_merge_in_place(d1, d2)

    wrapper_test_good(runnable=runnable)


@pytest.mark.parametrize(
    ("filename", "expected_text"),
    [
        ("split_run", "{n1}"),
        ("hyperlink_between_text", "{n1} toto {n2}"),
    ],
)
def test_merge_text(filename: str, expected_text: str):
    path_folder = PATH_TEST_DOCS_TESTSUITE / "docx/merge"
    path = path_folder / f"{filename}.docx"

    def runnable():
        d1 = Docx(path)

        # not in place
        ps = [normalize_runs(d1, p, inplace=False) for p in d1.paragraphs]
        text = "".join(p.text for p in ps)
        assert text == expected_text

        # check that the d1 did not changed
        text = "".join(p.text for p in d1.paragraphs)
        assert text == expected_text

        # in place
        for p in d1.paragraphs:
            normalize_runs(d1, p, inplace=True)
        text = "".join(p.text for p in d1.paragraphs)
        assert text == expected_text

    wrapper_test_good(runnable=runnable)


@pytest.mark.parametrize(
    ("filename1", "filename2"),
    [
        ("split_run_not", "split_run"),
    ],
)
def test_merge(filename1: str, filename2: str):
    path_folder = PATH_TEST_DOCS_TESTSUITE / "docx/merge"
    path1 = path_folder / f"{filename1}.docx"
    path2 = path_folder / f"{filename2}.docx"

    def runnable():
        d1 = Docx(path1)
        d2 = Docx(path2)

        _equals_merge_not_in_place(d1, to_merge=d2)
        _equals_merge_in_place(d1, merged=d2)

    wrapper_test_good(runnable=runnable)


# ------------------- Table -------------------


def _extract_first_table(path: Path) -> Tuple[Docx, DocxTable]:
    doc = Docx(path)
    return doc, DocxTable(doc, doc.tables[0])


@pytest.mark.parametrize(
    ("filename", "expected_dimensions"),
    [
        ("empty_rows2_cols3", (2, 3)),
        ("empty_rows3_cols2", (3, 2)),
        ("empty_merged_rows", (2, 3)),
        ("empty_merged_cols", (2, 3)),
    ],
)
def test_dimensions_table(filename: str, expected_dimensions: Tuple[int, int]):

    path = PATH_TEST_DOCS_TESTSUITE / f"docx/table/{filename}.docx"

    def runnable():
        _, table = _extract_first_table(path)
        assert table.get_dimensions() == expected_dimensions

    wrapper_test_good(runnable=runnable)


@pytest.mark.parametrize(
    ("filename", "row", "amount", "check_equal_docx"),
    [
        ("empty_rows2_cols3", 1, 1, False),
        ("empty_rows2_cols3", 2, 2, False),
        ("empty_rows2_cols3", 2, 2, False),
        ("empty_rows2_cols3", 3, 2, False),
        ("insert_rows_bold", 1, 1, True),
        ("insert_rows_bold", 2, 1, True),
        ("insert_rows_bold2", 1, 2, True),
    ],
)
def test_insert_rows2_table(
    filename: str, row: int, amount: int, check_equal_docx: bool
):

    path = PATH_TEST_DOCS_TESTSUITE / f"docx/table/{filename}.docx"
    path_output = PATH_TEST_DOCS_TESTSUITE / f"docx/table/{filename}_actual.docx"
    path_expected = PATH_TEST_DOCS_TESTSUITE / f"docx/table/{filename}_expected.docx"

    def runnable():
        doc, table = _extract_first_table(path)

        rows, cols = table.get_dimensions()
        table.insert_rows(row=row, amount=amount)

        doc.save(path_output)

        assert table.get_dimensions() == (rows + amount, cols)

        if check_equal_docx:
            assert docx_equals(doc, Docx(path_expected))

    wrapper_test_good(runnable=runnable)


@pytest.mark.parametrize(
    ("filename"),
    [
        "replace_text_in_cell_simple",
        "replace_text_in_cell_many_cells",
        "replace_text_in_cell_mix_style",
    ],
)
def test_replace_text_in_cell_table(filename: str):

    path = PATH_TEST_DOCS_TESTSUITE / f"docx/table/{filename}.docx"
    path_output = PATH_TEST_DOCS_TESTSUITE / f"docx/table/{filename}_actual.docx"
    path_expected = PATH_TEST_DOCS_TESTSUITE / f"docx/table/{filename}_expected.docx"

    def runnable():
        doc, table = _extract_first_table(path)
        for row in range(1, table.get_row_dimension() + 1):
            for col in range(1, table.get_col_dimension() + 1):
                table.replace_text_in_cell(
                    row=row,
                    col=col,
                    replace_text=build_replace_text({"n1": "v1", "n2": "v2"}),
                )

        doc.save(path_output)

        assert docx_equals(doc, Docx(path_expected))

    wrapper_test_good(runnable=runnable)


@pytest.mark.parametrize(
    ("filename"),
    ["copy_cell_simple", "copy_cell_bold", "copy_cell_mix_style"],
)
def test_copy_cell_table(
    filename: str,
):

    path = PATH_TEST_DOCS_TESTSUITE / f"docx/table/{filename}.docx"
    path_output = PATH_TEST_DOCS_TESTSUITE / f"docx/table/{filename}_actual.docx"
    path_expected = PATH_TEST_DOCS_TESTSUITE / f"docx/table/{filename}_expected.docx"

    def runnable():
        doc, table = _extract_first_table(path)
        table.copy_cell(src_cell=table.get_cell(1, 1), row=1, col=2)
        doc.save(path_output)

        assert docx_equals(doc, Docx(path_expected))

    wrapper_test_good(runnable=runnable)


@pytest.mark.parametrize(
    ("filename", "from_row", "from_col", "to_row", "to_col", "nb_row", "nb_col"),
    [
        ("copy_rectangle1", 1, 2, 2, 3, 2, 3),
        ("copy_rectangle2", 1, 2, 3, 2, 2, None),
    ],
)
def test_copy_rectanble_table(
    filename: str,
    from_row,
    from_col,
    to_row,
    to_col,
    nb_row,
    nb_col,
):

    path = PATH_TEST_DOCS_TESTSUITE / f"docx/table/{filename}.docx"
    path_output = PATH_TEST_DOCS_TESTSUITE / f"docx/table/{filename}_actual.docx"
    path_expected = PATH_TEST_DOCS_TESTSUITE / f"docx/table/{filename}_expected.docx"

    def runnable():
        doc, table = _extract_first_table(path)
        table.copy_rectangle(
            from_row=from_row,
            from_col=from_col,
            to_row=to_row,
            to_col=to_col,
            nb_row=nb_row,
            nb_col=nb_col,
        )
        doc.save(path_output)

        assert docx_equals(doc, Docx(path_expected))

    wrapper_test_good(runnable=runnable)


# @pytest.mark.parametrize(
#     ("filename", ""),
#     [],
# )
# def test__table(
#     filename: str,
# ):

#     path = PATH_TEST_DOCS_TESTSUITE / f"docx/table/{filename}.docx"

#     def runnable():
#         _, table = _extract_first_table(path)

#     wrapper_test_good(runnable=runnable)
